from docx import Document
import sys

try:
    doc = Document(r"e:\sga-ragflow-GM\新闻接口示例代码及返回案例\获取内网新闻信息\接口秘钥及附件地址说明.docx")
    print("=== 文档内容开始 ===")
    for para in doc.paragraphs:
        if para.text.strip():
            print(para.text)
    print("=== 文档内容结束 ===")

    print("\n=== 表格内容 ===")
    for table in doc.tables:
        for row in table.rows:
            print(" | ".join([cell.text.strip() for cell in row.cells]))
except Exception as e:
    print(f"读取失败: {e}")
